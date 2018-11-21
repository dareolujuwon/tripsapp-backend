import docx
import json


def convert_to_json(filename):
    """
    Convert from DOCX to JSON
    :param filename: .docx file
    :return: .json file
    """
    document = docx.Document(filename)
    document_to_json = []
    docs = []

    for row in range(0, len(document.paragraphs)):
        if row == 0:
            # Handle first row as header
            header = [x.lower() for x in document.paragraphs[row].text.split("\t") if x]
        else:
            # Handle the rest of the rows
            current_row = [x for x in document.paragraphs[row].text.split("\t") if x]
            # If column lengths don't match, we cant parse -> lets get out of here
            if len(header) != len(current_row):
                return print("Cannot parse document")
            document_to_json.append(dict(zip(header, current_row)))

    for a in range(0, len(document_to_json)):
        document_to_json[a]['primaryid'] = a+1
        docs.append(document_to_json[a])
    filepathnamewext = '' + filename[:-5] + '.json'
    with open(filepathnamewext, 'w') as fp:
        json.dump(docs, fp)
    return docs


def convert_to_docx(filename):
    """
    Convert from JSON to DOCX
    :param filename: .json file
    :return: .docx file
    """
    k = json.loads(open(filename).read())
    document = docx.Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'

    paragraph = document.add_paragraph()
    paragraph.add_run("{}\t{}\t{}\t{}\t{}\t{}".format("Id", "datetime", "description", "longitude", "latitude", "elevation"))
    for product in k:
        p = document.add_paragraph()
        for i in (0, len(product)):
            id = product['id']
            datetime = product['datetime']
            description = product['description']
            elevation = product['elevation']
            latitude = product['latitude']
            longitude = product['longitude']
            p.add_run(
                "{}\t{}\t{}\t{}\t{}\t{}".format(id, datetime, description, longitude, latitude, elevation))
            break
    # return document.save('input.txt.docx')
    document.save('/Users/olujuwondare/PycharmProjects/untitled1/files/input.txt.docx')


# reader = Reader()
# file = '/Users/olujuwondare/PycharmProjects/untitled1/files/input.txt.json'
# file = '/Users/olujuwondare/PycharmProjects/untitled1/files/input.txt.docx'
# a = reader.convert_to_json(file)
# b = reader.convert_to_docx(file)
